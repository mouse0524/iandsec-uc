from __future__ import annotations

import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree


SUPPORTED_MARKDOWN_UPLOAD_EXTENSIONS = {
    "pdf",
    "pptx",
    "docx",
    "xlsx",
    "html",
    "json",
    "txt",
    "md",
}

SUPPORTED_MARKDOWN_UPLOAD_MESSAGE = "仅支持 XLSX、PPTX、DOCX、PDF、HTML、JSON、TXT、MD 文件"


class SkillKnowDocumentParser:
    async def convert_to_markdown(self, file_path: str, file_type: str | None = None, **kwargs) -> str:
        ext = (file_type or Path(file_path).suffix).lower().lstrip(".")
        if ext not in SUPPORTED_MARKDOWN_UPLOAD_EXTENSIONS:
            raise ValueError(SUPPORTED_MARKDOWN_UPLOAD_MESSAGE)
        if ext in {"md", "txt"}:
            markdown = Path(file_path).read_text(encoding="utf-8", errors="ignore")
        else:
            try:
                from markitdown import MarkItDown

                result = MarkItDown().convert(file_path, **kwargs)
                markdown = result.text_content or ""
            except Exception as exc:
                fallback = self._fallback_to_markdown(file_path, ext)
                if fallback:
                    markdown = fallback
                else:
                    raise ValueError(f"Markdown 转换失败: {exc}") from exc
        markdown = self._normalize(markdown)
        if not markdown:
            raise ValueError("Markdown 转换结果为空")
        return markdown

    async def parse(self, file_path: str, file_type: str) -> str:
        return await self.convert_to_markdown(file_path, file_type)

    def _normalize(self, markdown: str) -> str:
        value = str(markdown or "").replace("\r\n", "\n").replace("\r", "\n")
        value = re.sub(r"\n{4,}", "\n\n\n", value)
        return value.strip()

    def _fallback_to_markdown(self, file_path: str, ext: str) -> str:
        try:
            if ext == "docx":
                return self._docx_to_markdown(file_path)
            if ext == "pptx":
                return self._pptx_to_markdown(file_path)
            if ext == "xlsx":
                return self._xlsx_to_markdown(file_path)
            if ext == "pdf":
                return self._pdf_to_markdown(file_path)
        except Exception:
            return ""
        return ""

    def _docx_to_markdown(self, file_path: str) -> str:
        from docx import Document

        document = Document(file_path)
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue
            width = max(len(row) for row in rows)
            normalized_rows = [row + [""] * (width - len(row)) for row in rows]
            parts.append("| " + " | ".join(normalized_rows[0]) + " |")
            parts.append("| " + " | ".join("---" for _ in range(width)) + " |")
            for row in normalized_rows[1:]:
                parts.append("| " + " | ".join(row) + " |")
        return "\n\n".join(parts)

    def _pptx_to_markdown(self, file_path: str) -> str:
        slides: list[tuple[int, str]] = []
        namespace = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        with zipfile.ZipFile(file_path) as archive:
            names = sorted(
                (name for name in archive.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)),
                key=lambda name: int(re.search(r"slide(\d+)\.xml$", name).group(1)),
            )
            for index, name in enumerate(names, start=1):
                root = ElementTree.fromstring(archive.read(name))
                texts = [node.text.strip() for node in root.findall(".//a:t", namespace) if node.text and node.text.strip()]
                if texts:
                    slides.append((index, "\n\n".join(texts)))
        return "\n\n".join(f"## 幻灯片 {index}\n\n{text}" for index, text in slides)

    def _xlsx_to_markdown(self, file_path: str) -> str:
        from openpyxl import load_workbook

        workbook = load_workbook(file_path, data_only=True, read_only=True)
        parts: list[str] = []
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if cell is None else str(cell).strip().replace("\n", " ") for cell in row]
                while cells and not cells[-1]:
                    cells.pop()
                if cells:
                    rows.append(cells)
            if not rows:
                continue
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            parts.append(f"## {sheet.title}")
            parts.append("| " + " | ".join(rows[0]) + " |")
            parts.append("| " + " | ".join("---" for _ in range(width)) + " |")
            for row in rows[1:]:
                parts.append("| " + " | ".join(row) + " |")
        return "\n\n".join(parts)

    def _pdf_to_markdown(self, file_path: str) -> str:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        parts: list[str] = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(f"## PDF Page {page_index}\n\n{text}")
        return "\n\n".join(parts)


skill_know_document_parser = SkillKnowDocumentParser()
