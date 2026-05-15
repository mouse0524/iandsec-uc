import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document

from app.services.skill_know.document_parser import SUPPORTED_MARKDOWN_UPLOAD_EXTENSIONS, skill_know_document_parser
from app.services.skill_know.markdown_chunker import skill_know_markdown_chunker


class SkillKnowMarkdownTestCase(unittest.TestCase):
    def test_supported_extensions_exclude_legacy_office_formats(self):
        for ext in ["xlsx", "pptx", "docx", "pdf", "html", "json", "txt", "md"]:
            self.assertIn(ext, SUPPORTED_MARKDOWN_UPLOAD_EXTENSIONS)

        for ext in ["doc", "ppt", "xls", "htm", "csv", "xml", "markdown"]:
            self.assertNotIn(ext, SUPPORTED_MARKDOWN_UPLOAD_EXTENSIONS)

    def test_chunker_preserves_heading_context(self):
        markdown = "# 产品手册\n\n## 登录认证\n\n" + "验证码错误时请重新获取验证码。\n\n" * 80

        chunks = skill_know_markdown_chunker.chunk(markdown, target_chars=300, max_chars=500, overlap_chars=50)

        self.assertGreater(len(chunks), 1)
        self.assertEqual(chunks[0].index, 0)
        self.assertTrue(any(chunk.heading and "产品手册" in chunk.heading for chunk in chunks))
        self.assertTrue(any(chunk.heading and "登录认证" in chunk.heading for chunk in chunks))
        self.assertTrue(all(chunk.content for chunk in chunks))
        self.assertTrue(all(chunk.token_count > 0 for chunk in chunks))

    def test_docx_falls_back_when_markitdown_dependency_is_missing(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "manual.docx"
            document = Document()
            document.add_paragraph("账号登录说明")
            document.add_paragraph("验证码错误时请重新获取。")
            document.save(path)

            with patch("markitdown.MarkItDown.convert", side_effect=Exception("DocxConverter MissingDependencyException")):
                markdown = self._run(skill_know_document_parser.convert_to_markdown(str(path), "docx"))

        self.assertIn("账号登录说明", markdown)
        self.assertIn("验证码错误时请重新获取", markdown)

    def test_pptx_falls_back_when_markitdown_image_extraction_fails(self):
        slide_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
  <p:cSld><p:spTree><p:sp><p:txBody>
    <a:p><a:r><a:t>数据安全方案</a:t></a:r></a:p>
    <a:p><a:r><a:t>AI 知识库辅助排障</a:t></a:r></a:p>
  </p:txBody></p:sp></p:spTree></p:cSld>
</p:sld>"""
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "deck.pptx"
            with zipfile.ZipFile(path, "w") as archive:
                archive.writestr("ppt/slides/slide1.xml", slide_xml)

            with patch("markitdown.MarkItDown.convert", side_effect=ValueError("no embedded image")):
                markdown = self._run(skill_know_document_parser.convert_to_markdown(str(path), "pptx"))

        self.assertIn("## 幻灯片 1", markdown)
        self.assertIn("数据安全方案", markdown)
        self.assertIn("AI 知识库辅助排障", markdown)

    def _run(self, awaitable):
        import asyncio

        return asyncio.run(awaitable)


if __name__ == "__main__":
    unittest.main()
