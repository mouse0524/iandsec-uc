from pathlib import Path
from datetime import datetime, timezone

import pytest
from fastapi import HTTPException
from openpyxl import Workbook
from PIL import Image

from app.api.v1.wiki import wiki as wiki_module
from app.services.wiki.markdown_converter import _save_docling_images, convert_to_markdown
from app.services.wiki import learning_service as learning_module
from app.services.wiki import import_service as import_module
from app.services.llm import LLMOpenAIClient
from app.services.wiki.import_service import WikiImportService
from app.services.wiki.learning_service import wiki_learning_service
from app.services.wiki.search_service import search_terms
from app.services.wiki.wiki_builder import content_hash, normalize_page_path, slugify, wiki_builder
from app.api.v1.wiki.wiki import _message_preview, _safe_child, _sse
from app.core.init_app import _old_ai_knowledge_menu_filter, _wiki_edit_api_paths, _wiki_read_api_paths


def test_wiki_text_and_csv_conversion(tmp_path: Path):
    text_path = tmp_path / "a.txt"
    text_path.write_text("hello wiki", encoding="utf-8")
    assert convert_to_markdown(text_path) == "hello wiki"

    csv_path = tmp_path / "a.csv"
    csv_path.write_text("name,value\nA,1\n", encoding="utf-8")
    markdown = convert_to_markdown(csv_path)
    assert "| name | value |" in markdown
    assert "| A | 1 |" in markdown


def test_wiki_xlsx_conversion(tmp_path: Path):
    path = tmp_path / "a.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Products"
    sheet.append(["name", "value"])
    sheet.append(["EDG", 2])
    workbook.save(path)

    markdown = convert_to_markdown(path)

    assert "## Products" in markdown
    assert "EDG" in markdown
    assert "2" in markdown


def test_wiki_pptx_is_not_supported(tmp_path: Path):
    with pytest.raises(ValueError):
        convert_to_markdown(tmp_path / "demo.pptx")


def test_wiki_docx_conversion_uses_docling(tmp_path: Path):
    from docx import Document

    image_path = tmp_path / "img.png"
    Image.new("RGB", (2, 2), color="red").save(image_path)
    docx_path = tmp_path / "with-image.docx"
    doc = Document()
    doc.add_paragraph("安装截图")
    doc.add_picture(str(image_path))
    doc.save(docx_path)

    markdown = convert_to_markdown(docx_path, asset_dir=tmp_path / "assets", asset_prefix="../assets")

    assert "安装截图" in markdown


def test_wiki_pdf_falls_back_to_page_images_when_docling_fails(monkeypatch, tmp_path: Path):
    from docling.document_converter import DocumentConverter

    pdf_path = tmp_path / "image.pdf"
    image_path = tmp_path / "image.png"
    Image.new("RGB", (20, 20), color="blue").save(image_path)
    Image.open(image_path).save(pdf_path)

    def fail_convert(self, path):
        raise RuntimeError("docling-parse could not load document")

    monkeypatch.setattr(DocumentConverter, "convert", fail_convert)
    markdown = convert_to_markdown(pdf_path, asset_dir=tmp_path / "assets", asset_prefix="assets/hash")

    assert "## Page 1" in markdown
    assert "assets/hash/image-page-1.png" in markdown
    assert (tmp_path / "assets" / "image-page-1.png").is_file()


def test_wiki_docling_images_are_written_to_assets(tmp_path: Path):
    class FakeDocument:
        images = {"media/image 1": b"png-bytes"}

    markdown = _save_docling_images(
        FakeDocument(),
        asset_dir=tmp_path / "assets",
        asset_prefix="assets/hash",
        markdown="# Demo",
    )

    assert "![image-1](assets/hash/image-1.png)" in markdown
    assert (tmp_path / "assets" / "image-1.png").read_bytes() == b"png-bytes"


@pytest.mark.anyio
async def test_wiki_source_upload_has_no_size_limit(monkeypatch, tmp_path: Path):
    saved = {}

    class FakeUpload:
        filename = "big.txt"
        chunks = [b"a", b"bc", b""]

        async def read(self, size=-1):
            return self.chunks.pop(0)

    class FakeQuery:
        async def first(self):
            return None

    class FakeSource:
        id = 7

        @staticmethod
        def filter(**kwargs):
            saved["filter"] = kwargs
            return FakeQuery()

        @staticmethod
        async def create(**kwargs):
            saved["source"] = kwargs
            return FakeSource()

    class FakeJob:
        @staticmethod
        async def create(**kwargs):
            saved["job"] = kwargs

    monkeypatch.setattr(import_module.settings, "MAX_UPLOAD_SIZE", 1)
    monkeypatch.setattr(import_module, "WikiSource", FakeSource)
    monkeypatch.setattr(import_module, "WikiIngestJob", FakeJob)
    service = WikiImportService()
    service.root = tmp_path

    source = await service.create_source(user_id=9, file=FakeUpload())

    assert source.id == 7
    assert saved["source"]["file_size"] == 3
    assert saved["job"] == {"source_id": 7}


@pytest.mark.anyio
async def test_wiki_source_can_be_created_from_chunked_file(monkeypatch, tmp_path: Path):
    saved = {}

    class FakeQuery:
        async def first(self):
            return None

    class FakeSource:
        id = 8

        @staticmethod
        def filter(**kwargs):
            saved["filter"] = kwargs
            return FakeQuery()

        @staticmethod
        async def create(**kwargs):
            saved["source"] = kwargs
            return FakeSource()

    class FakeJob:
        @staticmethod
        async def create(**kwargs):
            saved["job"] = kwargs

    monkeypatch.setattr(import_module, "WikiSource", FakeSource)
    monkeypatch.setattr(import_module, "WikiIngestJob", FakeJob)
    service = WikiImportService()
    service.root = tmp_path
    chunked = tmp_path / "uploading" / "merged.txt"
    chunked.parent.mkdir()
    chunked.write_bytes(b"abc")

    source = await service.create_source_from_path(user_id=9, filename="big.txt", raw_path=chunked)

    assert source.id == 8
    assert saved["source"]["file_size"] == 3
    assert saved["source"]["file_path"].endswith(".txt")
    assert saved["job"] == {"source_id": 8}
    assert not chunked.exists()


@pytest.mark.anyio
async def test_wiki_markdown_polish_accepts_safe_revision(monkeypatch, tmp_path: Path):
    class FakeLLM:
        async def chat(self, messages, timeout=None):
            return {"choices": [{"message": {"content": "```md\n# Demo\n\n整理后\n\n![图](assets/hash/a.png)\n```"}}]}

    class FakeSource:
        filename = "demo.txt"
        file_type = ".txt"

    raw = tmp_path / "demo.txt"
    raw.write_text("原文", encoding="utf-8")
    monkeypatch.setattr(import_module, "llm_openai_client", FakeLLM())

    result = await WikiImportService().polish_markdown(
        source=FakeSource(),
        raw_path=raw,
        markdown="# Demo\n\n原文\n\n![图](assets/hash/a.png)",
    )

    assert "整理后" in result


@pytest.mark.anyio
async def test_wiki_markdown_polish_rejects_missing_assets(monkeypatch, tmp_path: Path):
    class FakeLLM:
        async def chat(self, messages, timeout=None):
            return {"choices": [{"message": {"content": "# Demo\n\n丢图了"}}]}

    class FakeSource:
        filename = "demo.docx"
        file_type = ".docx"

    original = "# Demo\n\n原文\n\n![图](assets/hash/a.png)"
    monkeypatch.setattr(import_module, "llm_openai_client", FakeLLM())

    result = await WikiImportService().polish_markdown(
        source=FakeSource(),
        raw_path=tmp_path / "demo.docx",
        markdown=original,
    )

    assert result == original


def test_wiki_builder_falls_back_when_llm_json_is_invalid():
    payload = wiki_builder.normalize_llm_payload("not json", title="Install Guide", markdown="# Install")

    assert payload["source_page"]["title"] == "Install Guide"
    assert "## Source Notes" in payload["source_page"]["content"]
    assert "# Install" in payload["source_page"]["content"]
    assert payload["concepts"] == []


def test_wiki_builder_rejects_path_traversal(tmp_path: Path):
    builder = type(wiki_builder)()
    builder.wiki_dir = tmp_path

    with pytest.raises(ValueError):
        builder._write_page("../escape.md", "nope")


def test_wiki_safe_child_rejects_same_prefix_escape(tmp_path: Path):
    root = tmp_path / "wiki"
    root.mkdir()

    with pytest.raises(Exception):
        _safe_child(root, "../wiki_bak/escape.md")


def test_wiki_edit_permissions_include_chunk_upload_endpoints():
    paths = set(_wiki_edit_api_paths())

    assert "/api/v1/wiki/source/upload/init" in paths
    assert "/api/v1/wiki/source/upload/chunk" in paths
    assert "/api/v1/wiki/source/upload/complete" in paths


@pytest.mark.anyio
async def test_wiki_source_editor_requires_admin(monkeypatch):
    user = type("User", (), {"is_superuser": False})()

    async def fake_current_user():
        return user

    async def fake_role_names(_user):
        return ["技术"]

    monkeypatch.setattr(wiki_module, "_current_user", fake_current_user)
    monkeypatch.setattr(wiki_module, "_role_names", fake_role_names)

    with pytest.raises(HTTPException):
        await wiki_module._require_wiki_editor()


def test_wiki_read_permissions_include_view_source_endpoints():
    paths = set(_wiki_read_api_paths())

    assert "/api/v1/wiki/source/list" in paths
    assert "/api/v1/wiki/source/markdown" in paths


def test_old_ai_knowledge_menu_cleanup_targets_legacy_menu():
    def filters(q):
        data = dict(q.filters)
        for child in q.children:
            data.update(filters(child))
        return data

    data = filters(_old_ai_knowledge_menu_filter())

    assert data["name"] == "AI知识库"
    assert data["path"] == "/skill-know"
    assert data["component__startswith"] == "/skill-know"


def test_wiki_query_archive_markdown_contains_citations():
    content = wiki_builder._query_content(
        question='授权到期怎么续费',
        answer='按 [[续费流程]] 处理。',
        citations=[{"title": "续费流程", "path": "practices/renewal.md"}],
        created_at=datetime(2026, 7, 6, tzinfo=timezone.utc),
    )

    assert "type: query" in content
    assert "## Question\n授权到期怎么续费" in content
    assert "## Answer\n按 [[续费流程]] 处理。" in content
    assert "- [续费流程](practices/renewal.md)" in content


def test_wiki_paths_and_hashes_are_stable():
    assert slugify("License Renewal!") == "license-renewal"
    assert normalize_page_path("sources/monsafe常见问题解答-v1-0") == "sources/monsafe常见问题解答-v1-0.md"
    assert normalize_page_path(r"sources\demo.md") == "sources/demo.md"
    assert content_hash("same") == content_hash(b"same")


def test_wiki_sse_status_event_is_readable_json():
    event = _sse({"type": "status", "payload": {"stage": "search", "label": "查询 Wiki"}})

    assert event == 'data: {"type": "status", "payload": {"stage": "search", "label": "查询 Wiki"}}\n\n'


def test_wiki_message_preview_collapses_whitespace():
    assert _message_preview("  第一行\n第二行\t第三行  ", 8) == "第一行 第二行"


def test_wiki_search_terms_support_chinese_without_spaces():
    terms = search_terms("授权到期怎么续费")

    assert "授权" in terms
    assert "到期" in terms
    assert "续费" in terms
    assert "怎么" not in terms


def test_wiki_search_terms_prioritize_domain_dictionary():
    terms = search_terms("WebDAV同步失败如何处理")

    assert terms[:2] == ["webdav", "同步"]
    assert "如何" not in terms


def test_wiki_search_terms_accept_custom_dictionary():
    terms = search_terms("质保延期怎么办", domain_terms={"质保", "延期"}, stop_words={"怎么", "怎么办"})

    assert terms[:2] == ["质保", "延期"]
    assert "怎么办" not in terms


@pytest.mark.anyio
async def test_llm_client_reads_system_setting_config(monkeypatch):
    class FakeSettings:
        async def get_full_dict(self):
            return {
                "llm_chat_provider": "ollama",
                "llm_chat_base_url": "http://127.0.0.1:11434",
                "llm_chat_model": "qwen2.5",
                "llm_temperature": 0.1,
                "llm_timeout": 12,
            }

    import app.controllers.system_setting as settings_module

    monkeypatch.setattr(settings_module, "system_setting_controller", FakeSettings())

    config = await LLMOpenAIClient()._config()

    assert config["llm_chat_provider"] == "ollama"
    assert config["llm_chat_base_url"] == "http://127.0.0.1:11434"
    assert config["llm_chat_model"] == "qwen2.5"
    assert config["llm_timeout"] == 12


@pytest.mark.anyio
async def test_wiki_learning_candidate_reuses_pending(monkeypatch):
    saved = {}

    class FakeCandidate:
        answer = None
        evidence_page_ids = []

        async def save(self):
            saved["answer"] = self.answer
            saved["evidence_page_ids"] = self.evidence_page_ids

    class FakeQuery:
        async def first(self):
            return FakeCandidate()

    class FakeModel:
        @staticmethod
        def filter(**kwargs):
            saved["filter"] = kwargs
            return FakeQuery()

        @staticmethod
        async def create(**kwargs):
            raise AssertionError("should reuse pending candidate")

    monkeypatch.setattr(learning_module, "WikiLearningCandidate", FakeModel)

    candidate = await wiki_learning_service.create_candidate(
        question="授权到期怎么续费",
        answer="old answer",
        evidence_page_ids=[1, 2],
        reason="unhelpful_feedback",
    )

    assert isinstance(candidate, FakeCandidate)
    assert saved["filter"]["status"] == "pending"
    assert saved["answer"] == "old answer"
    assert saved["evidence_page_ids"] == [1, 2]
